"""PDF to DOCX conversion service."""

from __future__ import annotations

from io import BytesIO
from pathlib import Path
from tempfile import TemporaryDirectory

import pdfplumber
import pypdfium2 as pdfium
from pdf2docx import Converter as Pdf2DocxConverter
from docx import Document
from docx.shared import Inches

from ..core.exceptions import ConversionError


def _render_page_as_png(pdf_document: pdfium.PdfDocument, page_index: int) -> bytes:
    """Render a PDF page to PNG bytes for image-preserving DOCX output."""
    page = pdf_document[page_index]
    bitmap = page.render(scale=1.8)
    pil_image = bitmap.to_pil()
    png_stream = BytesIO()
    pil_image.save(png_stream, format="PNG")
    page.close()
    pdf_document.close()
    png_stream.seek(0)
    return png_stream.read()


def _convert_with_pdf2docx(pdf_bytes: bytes) -> bytes | None:
    """Use pdf2docx for higher-fidelity conversion with image and layout support."""
    with TemporaryDirectory() as tmp_dir:
        tmp_path = Path(tmp_dir)
        input_path = tmp_path / "input.pdf"
        output_path = tmp_path / "output.docx"
        input_path.write_bytes(pdf_bytes)

        converter = Pdf2DocxConverter(str(input_path))
        try:
            converter.convert(str(output_path))
        finally:
            converter.close()

        if output_path.exists() and output_path.stat().st_size > 0:
            return output_path.read_bytes()

    return None


def convert_pdf_to_docx(pdf_bytes: bytes) -> bytes:
    """Extract text and preserve visuals by embedding rendered page previews."""
    pdfium_document = None
    try:
        # Preferred path: pdf2docx generally keeps text, images, and layout better.
        try:
            converted = _convert_with_pdf2docx(pdf_bytes)
            if converted:
                return converted
        except Exception:
            pass

        # Fallback path: text extraction + page previews to ensure images are still visible.
        page_sections: list[tuple[str, bytes | None]] = []
        pdfium_document = pdfium.PdfDocument(pdf_bytes)

        with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
            for index, page in enumerate(pdf.pages, start=1):
                extracted = page.extract_text() or ""
                section_header = f"--- Page {index} ---"
                text_block = f"{section_header}\n{extracted}".strip()

                # Always include a page preview to preserve images and layout-like context.
                page_image_bytes: bytes | None = None
                try:
                    page_image_bytes = _render_page_as_png(pdfium_document, index - 1)
                except Exception:
                    page_image_bytes = None

                page_sections.append((text_block, page_image_bytes))

        document = Document()
        for block, page_image in page_sections:
            document.add_paragraph(block)
            if page_image:
                document.add_picture(BytesIO(page_image), width=Inches(6.2))
            document.add_paragraph("")

        output_stream = BytesIO()
        document.save(output_stream)
        output_stream.seek(0)
        return output_stream.read()
    except Exception as exc:
        raise ConversionError("Unable to convert PDF to DOCX.") from exc
    finally:
        if pdfium_document is not None:
            pdfium_document.close()
